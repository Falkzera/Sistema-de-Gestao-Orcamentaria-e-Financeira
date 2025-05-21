from weasyprint import HTML
import markdown
import tempfile

def gerar_pdf_weasy_padrao(conteudo_pdf, nome_arquivo="relatorio.pdf"):
    # Juntar o conteúdo e converter markdown → HTML
    html_conteudo = "\n".join(conteudo_pdf)  # conteúdo já é HTML, não precisa de conversão
    html_final = f"""
    <html>
    <head>
        <meta charset="utf-8">
        <style>

            /* Define a página */
            @page {{
                size: A4;
                margin: 3cm 2cm 2cm 3cm;  /* topo, direita, inferior, esquerda */
                
                @bottom-right {{
                    font-family: 'Times New Roman', serif;
                    content: "Página " counter(page);
                    font-size: 10pt;
                }}
            }}

            body {{
                font-family: 'Times New Roman', serif;
                font-size: 12pt;
                line-height: 1.2;
                color: #000;
                text-align: justify;
            }}

            h1 {{
                text-align: center;
                font-size: 20pt;
                font-weight: bold;
                margin-top: 0;
                margin-bottom: 20px;
            }}

            h2 {{
                font-size: 16pt;
                font-weight: bold;
                margin-top: 30px;
                margin-bottom: 15px;
            }}

            h3 {{
                font-size: 14pt;
                font-weight: bold;
                margin-top: 20px;
                margin-bottom: 10px;
            }}

            p {{
                text-indent: 1.25cm;
                margin: 0 0 12pt 0;
            }}

            strong {{
                font-weight: bold;
            }}

            em {{
                font-style: italic;
            }}

        </style>
    </head>
    <body>
        {html_conteudo}
    </body>
    </html>
    """

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
        HTML(string=html_final).write_pdf(f.name)
        return f.name




from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import tempfile

def gerar_docx(conteudo_pdf, nome_arquivo="relatorio.docx"):
    """
    Gera um arquivo DOCX a partir do conteúdo (que é uma lista de strings HTML).

    Args:
        conteudo_pdf (list): Lista de strings (com HTML simples ou texto plano).
        nome_arquivo (str): Nome do arquivo para salvar.

    Returns:
        str: Caminho do arquivo gerado.
    """
    document = Document()

    # Definir fonte padrão
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Tratamento de cada linha
    for item in conteudo_pdf:
        if item.strip() == "":
            continue  # ignora linhas vazias

        # Interpretar tags simples
        if "<h1>" in item:
            text = re.sub(r'<\/?h1>', '', item)
            p = document.add_heading(text.strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "<h2>" in item:
            text = re.sub(r'<\/?h2>', '', item)
            document.add_heading(text.strip(), level=2)
        elif "<h3>" in item:
            text = re.sub(r'<\/?h3>', '', item)
            document.add_heading(text.strip(), level=3)
        elif "<p>" in item:
            text = re.sub(r'<\/?p>', '', item)
            paragraph = document.add_paragraph(text.strip())
            paragraph.paragraph_format.first_line_indent = Inches(0.5)  # 1.25 cm
            paragraph.paragraph_format.space_after = Pt(12)
        elif "<strong>" in item or "<em>" in item:
            paragraph = document.add_paragraph()
            run = paragraph.add_run()

            # Negrito e itálico
            text = item
            if "<strong>" in text:
                bold_text = re.findall(r'<strong>(.*?)<\/strong>', text)
                for part in bold_text:
                    r = paragraph.add_run(part)
                    r.bold = True
                    text = text.replace(f"<strong>{part}</strong>", "")
            if "<em>" in text:
                italic_text = re.findall(r'<em>(.*?)<\/em>', text)
                for part in italic_text:
                    r = paragraph.add_run(part)
                    r.italic = True
                    text = text.replace(f"<em>{part}</em>", "")
            
            clean_text = re.sub(r'<[^>]+>', '', text)  # remove outras tags
            paragraph.add_run(clean_text)
        else:
            # Texto comum
            paragraph = document.add_paragraph(item.strip())
            paragraph.paragraph_format.first_line_indent = Inches(0.5)
            paragraph.paragraph_format.space_after = Pt(12)

    # Salvar em arquivo temporário
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
        document.save(f.name)
        return f.name
